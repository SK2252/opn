"""
Validation Agent for the Document Agent System.

Provides comprehensive validation of input data including:
- File existence and format checking
- Column validation for Excel files
- Data type and quality checks
- Template validation for Word documents
"""

import os
from typing import Any, Dict, List, Optional, Set

import pandas as pd

from AI_open_negotiation.agents.document_agent.base_agent import BaseAgent
from AI_open_negotiation.models.task_models import DocumentTask, DocumentType, TaskStatus
from AI_open_negotiation.models.result_models import ValidationResult


class ValidationAgent(BaseAgent):
    """
    Agent responsible for validating input data before processing.
    
    Validates:
    - Excel file existence and format
    - Required columns presence
    - Data types and formats
    - Word template existence and placeholders
    - Data quality checks (null values, duplicates)
    
    Example:
        >>> agent = ValidationAgent("Validator", {"log_level": "DEBUG"})
        >>> task = DocumentTask(
        ...     task_id="val_001",
        ...     document_type=DocumentType.OPEN_NEG_GROUP,
        ...     input_data={"excel_path": "data.xlsx"}
        ... )
        >>> result = await agent.execute(task)
    """
    
    # Required columns for Open Negotiation documents
    REQUIRED_COLUMNS_GROUP: Set[str] = {
        "ProvOrgNPI",
        "Provider",
        "InsurancePlanName",
        "OpenNegGroup",
        "CPT_Description",
        "Claim Number",
        "Date of item(s) or service(s)",
        "Service code(s)",
        "Initial Payment",
        "Offer",
    }
    
    REQUIRED_COLUMNS_NOTICE: Set[str] = {
        "ProvOrgNPI",
        "Hospital Name",
        "Provider",
        "InsurancePlanName",
        "OpenNegNotice",
        "Notice Date",
        "CMS Date1",
        "CMS Date2",
    }
    
    # Required template placeholders
    TEMPLATE_PLACEHOLDERS: Set[str] = {
        "{Hospital Name}",
        "{Provider}",
        "{InsurancePlanName}",
        "{Notice Date}",
        "{CMS Date1}",
        "{CMS Date2}",
    }
    
    def __init__(self, name: str = "ValidationAgent", config: Optional[Dict[str, Any]] = None):
        """Initialize the ValidationAgent."""
        super().__init__(name, config)
    
    async def execute(self, task: DocumentTask) -> DocumentTask:
        """
        Execute validation on the input data.
        
        Args:
            task: Document task containing input data paths
            
        Returns:
            Task with validation result in metadata
        """
        self.start_timer()
        task.mark_in_progress()
        self.log_info(f"Starting validation for task {task.task_id}")
        
        try:
            input_data = task.input_data
            validation_result = ValidationResult(is_valid=True)
            
            # Validate Excel file
            excel_path = input_data.get("excel_path")
            if excel_path:
                excel_result = await self._validate_excel(
                    excel_path, 
                    task.document_type
                )
                if not excel_result.is_valid:
                    validation_result.is_valid = False
                validation_result.errors.extend(excel_result.errors)
                validation_result.warnings.extend(excel_result.warnings)
                validation_result.total_records = excel_result.total_records
                validation_result.validated_records = excel_result.validated_records
            
            # Validate template if provided
            template_path = input_data.get("template_docx")
            if template_path:
                template_result = await self._validate_template(template_path)
                if not template_result.is_valid:
                    validation_result.is_valid = False
                validation_result.errors.extend(template_result.errors)
                validation_result.warnings.extend(template_result.warnings)
            
            # Validate output folders
            for folder_key in ["output_group_folder", "output_notice_folder", "merged_output_folder"]:
                folder_path = input_data.get(folder_key)
                if folder_path:
                    folder_result = self._validate_output_folder(folder_path, folder_key)
                    validation_result.warnings.extend(folder_result.warnings)
            
            # Store result in task metadata
            task.metadata["validation_result"] = validation_result.to_dict()
            
            if validation_result.is_valid:
                task.mark_completed({"validation": validation_result.to_dict()})
                self.log_info(
                    f"Validation passed: {validation_result.validated_records}/"
                    f"{validation_result.total_records} records valid"
                )
            else:
                task.mark_failed(f"Validation failed: {validation_result.errors}")
                self.log_error(f"Validation failed with {len(validation_result.errors)} errors")
            
            if validation_result.warnings:
                self.log_warning(f"Validation warnings: {validation_result.warnings}")
            
            return task
            
        except Exception as e:
            self.log_error(f"Validation error: {e}")
            task.mark_failed(str(e))
            return task
    
    async def _validate_excel(
        self, 
        excel_path: str, 
        doc_type: DocumentType
    ) -> ValidationResult:
        """
        Validate Excel file structure and content.
        
        Args:
            excel_path: Path to Excel file
            doc_type: Document type to determine required columns
            
        Returns:
            ValidationResult with errors and warnings
        """
        result = ValidationResult(is_valid=True)
        
        # Check file existence
        if not os.path.exists(excel_path):
            result.add_error(f"Excel file not found: {excel_path}")
            return result
        
        # Check file extension
        if not excel_path.lower().endswith(('.xls', '.xlsx')):
            result.add_error(f"Invalid file format: {excel_path}. Expected .xls or .xlsx")
            return result
        
        try:
            # Read Excel file
            df = pd.read_excel(excel_path)
            df.columns = df.columns.str.strip()
            result.total_records = len(df)
            
            # Determine required columns based on document type
            if doc_type == DocumentType.OPEN_NEG_GROUP:
                required_cols = self.REQUIRED_COLUMNS_GROUP
            elif doc_type == DocumentType.OPEN_NEG_NOTICE:
                required_cols = self.REQUIRED_COLUMNS_NOTICE
            else:
                # For general processing, require union of both
                required_cols = self.REQUIRED_COLUMNS_GROUP.union(self.REQUIRED_COLUMNS_NOTICE)
            
            # Check for missing columns
            actual_cols = set(df.columns)
            missing_cols = required_cols - actual_cols
            if missing_cols:
                result.add_error(f"Missing required columns: {missing_cols}")
            
            # Data quality checks
            if "ProvOrgNPI" in df.columns:
                null_npi = df["ProvOrgNPI"].isna().sum()
                if null_npi > 0:
                    result.add_warning(f"{null_npi} rows with missing ProvOrgNPI")
            
            if "InsurancePlanName" in df.columns:
                null_plan = df["InsurancePlanName"].isna().sum()
                if null_plan > 0:
                    result.add_warning(f"{null_plan} rows with missing InsurancePlanName")
            
            # Check for OpenNegGroup or OpenNegNotice values
            if doc_type == DocumentType.OPEN_NEG_GROUP and "OpenNegGroup" in df.columns:
                valid_records = int(df["OpenNegGroup"].notna().sum())
                result.validated_records = valid_records
                if valid_records == 0:
                    result.add_error("No valid OpenNegGroup values found")
            elif doc_type == DocumentType.OPEN_NEG_NOTICE and "OpenNegNotice" in df.columns:
                valid_records = int(df["OpenNegNotice"].notna().sum())
                result.validated_records = valid_records
                if valid_records == 0:
                    result.add_error("No valid OpenNegNotice values found")
            else:
                result.validated_records = result.total_records
            
            # Check for duplicates
            if "Claim Number" in df.columns:
                duplicates = df["Claim Number"].duplicated().sum()
                if duplicates > 0:
                    result.add_warning(f"{duplicates} duplicate claim numbers found")
            
            self.log_debug(f"Excel validation: {result.total_records} rows, {len(df.columns)} columns")
            
        except Exception as e:
            result.add_error(f"Failed to read Excel file: {str(e)}")
        
        return result
    
    async def _validate_template(self, template_path: str) -> ValidationResult:
        """
        Validate Word template file.
        
        Args:
            template_path: Path to Word template
            
        Returns:
            ValidationResult with errors and warnings
        """
        result = ValidationResult(is_valid=True)
        
        # Check file existence
        if not os.path.exists(template_path):
            result.add_error(f"Template file not found: {template_path}")
            return result
        
        # Check file extension
        if not template_path.lower().endswith(('.doc', '.docx')):
            result.add_error(f"Invalid template format: {template_path}. Expected .doc or .docx")
            return result
        
        try:
            from docx import Document
            
            doc = Document(template_path)
            
            # Extract all text from document
            all_text = ""
            for paragraph in doc.paragraphs:
                all_text += paragraph.text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text += cell.text
            
            # Check for placeholders
            missing_placeholders = []
            for placeholder in self.TEMPLATE_PLACEHOLDERS:
                if placeholder not in all_text:
                    missing_placeholders.append(placeholder)
            
            if missing_placeholders:
                result.add_warning(f"Missing template placeholders: {missing_placeholders}")
            
            self.log_debug(f"Template validation complete: {len(missing_placeholders)} missing placeholders")
            
        except Exception as e:
            result.add_error(f"Failed to read template file: {str(e)}")
        
        return result
    
    def _validate_output_folder(self, folder_path: str, folder_name: str) -> ValidationResult:
        """
        Validate output folder accessibility.
        
        Args:
            folder_path: Path to output folder
            folder_name: Name of the folder configuration key
            
        Returns:
            ValidationResult with warnings
        """
        result = ValidationResult(is_valid=True)
        
        if os.path.exists(folder_path):
            # Check if folder is writable
            if not os.access(folder_path, os.W_OK):
                result.add_warning(f"Output folder {folder_name} may not be writable: {folder_path}")
            
            # Check for existing files
            existing_files = len([f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))])
            if existing_files > 0:
                result.add_warning(f"Output folder {folder_name} contains {existing_files} existing files")
        else:
            # Folder will be created during processing
            self.log_debug(f"Output folder {folder_name} will be created: {folder_path}")
        
        return result
    
    async def validate_data(
        self,
        excel_path: str,
        template_docx: Optional[str] = None,
        document_type: DocumentType = DocumentType.OPEN_NEG_GROUP
    ) -> ValidationResult:
        """
        Convenience method for direct validation without task wrapper.
        
        Args:
            excel_path: Path to Excel file
            template_docx: Optional path to Word template
            document_type: Type of document to validate for
            
        Returns:
            ValidationResult with all validation results
        """
        task = DocumentTask(
            task_id="direct_validation",
            document_type=document_type,
            input_data={
                "excel_path": excel_path,
                "template_docx": template_docx,
            }
        )
        
        result_task = await self.execute(task)
        return ValidationResult(**result_task.metadata.get("validation_result", {"is_valid": False}))
