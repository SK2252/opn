"""
Generation Agents for the Document Agent System.

Contains specialized agents for generating different document types:
- GroupGenerationAgent: Generates Open Negotiation Group Excel files
- NoticeGenerationAgent: Generates Open Negotiation Notice Word/PDF files
"""

import os
import time
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd

from AI_open_negotiation.agents.document_agent.base_agent import BaseAgent
from AI_open_negotiation.models.task_models import DocumentTask, DocumentType, TaskStatus
from AI_open_negotiation.models.result_models import GenerationStats


class GroupGenerationAgent(BaseAgent):
    """
    Agent for generating Open Negotiation Group Excel files.
    
    Generates Excel files grouped by ProvOrgNPI and InsurancePlanName,
    with formatted columns for claim data.
    
    Example:
        >>> agent = GroupGenerationAgent("GroupGenerator", {})
        >>> task = DocumentTask(
        ...     task_id="grp_001",
        ...     document_type=DocumentType.OPEN_NEG_GROUP,
        ...     input_data={
        ...         "excel_path": "data.xlsx",
        ...         "output_group_folder": "output/groups"
        ...     }
        ... )
        >>> result = await agent.execute(task)
    """
    
    # Column mapping for output Excel
    COLUMN_MAPPING = {
        'CPT_Description': 'Description of item(s) and/or service(s)',
        'Claim Number': 'Claim Number',
        'ProvOrgNPI': 'Name of provider, facility, or provider of air ambulance services, and National Provider Identifier (NPI)',
        'Date of item(s) or service(s)': 'Date provided',
        'Service code(s)': 'Service code',
        'Initial Payment': 'Initial payment (if no initial payment amount, write N/A)',
        'Offer': 'Offer for total out-of- network rate (including any cost sharing)'
    }
    
    def __init__(self, name: str = "GroupGenerationAgent", config: Optional[Dict[str, Any]] = None):
        """Initialize the GroupGenerationAgent."""
        super().__init__(name, config)
    
    async def execute(self, task: DocumentTask) -> DocumentTask:
        """
        Execute group document generation.
        
        Args:
            task: Document task with excel_path and output_group_folder
            
        Returns:
            Task with generation statistics in metadata
        """
        self.start_timer()
        task.mark_in_progress()
        self.log_info(f"Starting group generation for task {task.task_id}")
        
        try:
            input_data = task.input_data
            excel_path = input_data.get("excel_path")
            output_folder = input_data.get("output_group_folder")
            
            if not excel_path or not output_folder:
                task.mark_failed("Missing excel_path or output_group_folder in input_data")
                return task
            
            stats = await self._generate_groups(excel_path, output_folder)
            stats.duration_seconds = self.get_elapsed_seconds()
            
            task.metadata["stats"] = stats.to_dict()
            
            if stats.failed == 0:
                task.mark_completed(stats.to_dict())
                self.log_info(f"Group generation completed: {stats.successful} files created")
            else:
                task.mark_partial_failure(
                    f"{stats.failed} groups failed to generate",
                    stats.to_dict()
                )
                self.log_warning(f"Partial failure: {stats.successful} success, {stats.failed} failed")
            
            return task
            
        except Exception as e:
            self.log_error(f"Group generation error: {e}")
            task.mark_failed(str(e))
            return task
    
    async def _generate_groups(self, excel_path: str, output_folder: str) -> GenerationStats:
        """
        Generate group Excel files from input data.
        
        Args:
            excel_path: Path to input Excel file
            output_folder: Base folder for output files
            
        Returns:
            GenerationStats with counts and file paths
        """
        stats = GenerationStats()
        
        # Read and prepare data
        df = pd.read_excel(excel_path)
        df.columns = df.columns.str.strip()
        os.makedirs(output_folder, exist_ok=True)
        
        # Get unique combinations
        unique_combinations = df[['ProvOrgNPI', 'Provider', 'InsurancePlanName']].drop_duplicates()
        stats.total_records = len(unique_combinations)
        
        self.log_info(f"Processing {stats.total_records} unique NPI/Insurance combinations")
        
        for idx, row in unique_combinations.iterrows():
            try:
                # Filter data for this combination
                filtered_df = df[
                    (df['ProvOrgNPI'] == row['ProvOrgNPI']) &
                    (df['Provider'] == row['Provider']) &
                    (df['InsurancePlanName'] == row['InsurancePlanName'])
                ]
                
                # Skip if no OpenNegGroup value
                if filtered_df.empty or pd.isna(filtered_df.iloc[0].get('OpenNegGroup')):
                    stats.add_skipped()
                    continue
                
                # Create output DataFrame with renamed columns
                output_df = filtered_df[list(self.COLUMN_MAPPING.keys())].rename(
                    columns=self.COLUMN_MAPPING
                )
                
                # Format date column
                output_df['Date provided'] = pd.to_datetime(
                    output_df['Date provided'], 
                    errors='coerce'
                ).dt.strftime('%b %d, %Y')
                
                # Add serial number column
                output_df.insert(0, 'SNO', range(1, len(output_df) + 1))
                
                # Format currency columns
                output_df['Initial payment (if no initial payment amount, write N/A)'] = \
                    output_df['Initial payment (if no initial payment amount, write N/A)'].apply(
                        self._format_currency
                    )
                output_df['Offer for total out-of- network rate (including any cost sharing)'] = \
                    output_df['Offer for total out-of- network rate (including any cost sharing)'].apply(
                        self._format_currency
                    )
                
                # Create output path
                safe_npi = self._safe_filename(str(row['ProvOrgNPI']))
                insurance = str(row['InsurancePlanName']).strip()
                output_path = os.path.join(output_folder, safe_npi, insurance)
                os.makedirs(output_path, exist_ok=True)
                
                # Get filename from data
                group_filename = str(filtered_df.iloc[0]['OpenNegGroup']).strip()
                if not group_filename.lower().endswith('.xlsx'):
                    group_filename += '.xlsx'
                
                # Save Excel file
                full_path = os.path.join(output_path, group_filename)
                output_df.to_excel(full_path, index=False)
                
                stats.add_success(full_path)
                self.log_debug(f"Created: {full_path}")
                
            except Exception as e:
                stats.add_failure(f"Failed for NPI {row.get('ProvOrgNPI')}: {str(e)}")
                self.log_error(f"Error processing group: {e}")
        
        return stats
    
    @staticmethod
    def _format_currency(x) -> str:
        """Format value as currency or N/A."""
        if pd.isna(x) or str(x).strip().upper() == 'N/A':
            return 'N/A'
        try:
            value = float(str(x).replace('$', '').replace(',', '').strip())
            return f"${value:,.2f}"
        except (ValueError, TypeError):
            return str(x)
    
    @staticmethod
    def _safe_filename(value: str) -> str:
        """Convert value to safe filename."""
        return "".join(c if c.isalnum() else "_" for c in value)


class NoticeGenerationAgent(BaseAgent):
    """
    Agent for generating Open Negotiation Notice Word/PDF files.
    
    Generates Word documents from templates with placeholder replacement,
    and converts them to PDF format.
    
    Example:
        >>> agent = NoticeGenerationAgent("NoticeGenerator", {})
        >>> task = DocumentTask(
        ...     task_id="ntc_001",
        ...     document_type=DocumentType.OPEN_NEG_NOTICE,
        ...     input_data={
        ...         "excel_path": "data.xlsx",
        ...         "template_docx": "template.docx",
        ...         "output_notice_folder": "output/notices"
        ...     }
        ... )
        >>> result = await agent.execute(task)
    """
    
    def __init__(self, name: str = "NoticeGenerationAgent", config: Optional[Dict[str, Any]] = None):
        """Initialize the NoticeGenerationAgent."""
        super().__init__(name, config)
    
    async def execute(self, task: DocumentTask) -> DocumentTask:
        """
        Execute notice document generation.
        
        Args:
            task: Document task with excel_path, template_docx, and output_notice_folder
            
        Returns:
            Task with generation statistics in metadata
        """
        self.start_timer()
        task.mark_in_progress()
        self.log_info(f"Starting notice generation for task {task.task_id}")
        
        try:
            input_data = task.input_data
            excel_path = input_data.get("excel_path")
            template_path = input_data.get("template_docx")
            output_folder = input_data.get("output_notice_folder")
            
            if not all([excel_path, template_path, output_folder]):
                task.mark_failed("Missing required input_data fields")
                return task
            
            stats = await self._generate_notices(excel_path, template_path, output_folder)
            stats.duration_seconds = self.get_elapsed_seconds()
            
            task.metadata["stats"] = stats.to_dict()
            
            if stats.failed == 0:
                task.mark_completed(stats.to_dict())
                self.log_info(f"Notice generation completed: {stats.successful} files created")
            else:
                task.mark_partial_failure(
                    f"{stats.failed} notices failed to generate",
                    stats.to_dict()
                )
                self.log_warning(f"Partial failure: {stats.successful} success, {stats.failed} failed")
            
            return task
            
        except Exception as e:
            self.log_error(f"Notice generation error: {e}")
            task.mark_failed(str(e))
            return task
    
    async def _generate_notices(
        self, 
        excel_path: str, 
        template_path: str, 
        output_folder: str
    ) -> GenerationStats:
        """
        Generate notice Word/PDF files from input data.
        
        Args:
            excel_path: Path to input Excel file
            template_path: Path to Word template
            output_folder: Base folder for output files
            
        Returns:
            GenerationStats with counts and file paths
        """
        from docx import Document
        import pythoncom
        
        stats = GenerationStats()
        
        # Initialize COM for PDF conversion
        pythoncom.CoInitialize()
        
        try:
            # Read and prepare data
            df = pd.read_excel(excel_path, engine='openpyxl')
            df.columns = df.columns.str.strip()
            
            # Get unique combinations for notices
            df_unique = df.drop_duplicates(
                subset=['ProvOrgNPI', 'Hospital Name', 'OpenNegNotice', 'InsurancePlanName']
            )
            stats.total_records = len(df_unique)
            
            os.makedirs(output_folder, exist_ok=True)
            self.log_info(f"Processing {stats.total_records} unique notice combinations")
            
            for idx, row in df_unique.iterrows():
                try:
                    # Skip if no OpenNegNotice value
                    if pd.isna(row.get('OpenNegNotice')):
                        stats.add_skipped()
                        continue
                    
                    # Load template
                    doc = Document(template_path)
                    
                    # Prepare replacements
                    replacements = {
                        '{Hospital Name}': str(row.get('Hospital Name', '')),
                        '{Provider}': str(row.get('Provider', '')),
                        '{InsurancePlanName}': str(row.get('InsurancePlanName', '')),
                        '{Notice Date}': str(row.get('Notice Date', '')),
                        '{CMS Date1}': str(row.get('CMS Date1', '')),
                        '{CMS Date2}': str(row.get('CMS Date2', '')),
                    }
                    bold_keys = set(replacements.keys())
                    
                    # Replace placeholders in paragraphs
                    for paragraph in doc.paragraphs:
                        self._replace_placeholders(paragraph, replacements, bold_keys)
                    
                    # Replace placeholders in tables
                    for table in doc.tables:
                        for table_row in table.rows:
                            for cell in table_row.cells:
                                for paragraph in cell.paragraphs:
                                    self._replace_placeholders(paragraph, replacements, bold_keys)
                    
                    # Create output path
                    save_path = os.path.join(
                        output_folder, 
                        str(row['ProvOrgNPI']), 
                        str(row['InsurancePlanName'])
                    )
                    os.makedirs(save_path, exist_ok=True)
                    
                    # Prepare filenames
                    notice_filename = str(row['OpenNegNotice']).strip()
                    base_filename = os.path.splitext(notice_filename)[0]
                    docx_path = os.path.join(save_path, base_filename + ".docx")
                    pdf_path = os.path.join(save_path, base_filename + ".pdf")
                    
                    # Save Word document
                    doc.save(docx_path)
                    
                    # Convert to PDF
                    try:
                        from docx2pdf import convert
                        convert(docx_path, pdf_path)
                        stats.add_success(pdf_path)
                    except Exception as pdf_error:
                        # PDF conversion failed, but Word doc was saved
                        self.log_warning(f"PDF conversion failed: {pdf_error}")
                        stats.add_success(docx_path)
                    
                    self.log_debug(f"Created: {docx_path}")
                    
                except Exception as e:
                    stats.add_failure(f"Failed for NPI {row.get('ProvOrgNPI')}: {str(e)}")
                    self.log_error(f"Error processing notice: {e}")
            
            return stats
            
        finally:
            pythoncom.CoUninitialize()
    
    def _replace_placeholders(
        self, 
        paragraph, 
        replacements: Dict[str, str], 
        bold_keys: set
    ) -> None:
        """
        Replace placeholders in a paragraph with values.
        
        Args:
            paragraph: docx Paragraph object
            replacements: Dictionary of placeholder -> replacement value
            bold_keys: Set of placeholders that should be bolded
        """
        full_text = "".join(run.text for run in paragraph.runs)
        
        replaced = False
        for placeholder, replacement in replacements.items():
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, replacement)
                replaced = True
        
        if replaced:
            # Clear existing runs
            for run in paragraph.runs:
                run.text = ""
            
            # Rebuild with formatting
            i = 0
            while i < len(full_text):
                match_found = False
                for placeholder, replacement in replacements.items():
                    rep_len = len(replacement)
                    if full_text[i:i + rep_len] == replacement:
                        run = paragraph.add_run(replacement)
                        if placeholder in bold_keys:
                            run.bold = True
                        i += rep_len
                        match_found = True
                        break
                
                if not match_found:
                    paragraph.add_run(full_text[i])
                    i += 1


class MergeAgent(BaseAgent):
    """
    Agent for merging output folders from different generation agents.
    
    Combines files from group and notice folders into a single merged output,
    with automatic renaming to avoid conflicts.
    """
    
    ALLOWED_EXTENSIONS = {".xls", ".xlsx", ".doc", ".docx", ".pdf"}
    
    def __init__(self, name: str = "MergeAgent", config: Optional[Dict[str, Any]] = None):
        """Initialize the MergeAgent."""
        super().__init__(name, config)
    
    async def execute(self, task: DocumentTask) -> DocumentTask:
        """
        Execute folder merge operation.
        
        Args:
            task: Document task with folder paths
            
        Returns:
            Task with merge statistics
        """
        self.start_timer()
        task.mark_in_progress()
        self.log_info(f"Starting merge for task {task.task_id}")
        
        try:
            input_data = task.input_data
            folder1 = input_data.get("output_group_folder")
            folder2 = input_data.get("output_notice_folder")
            output_folder = input_data.get("merged_output_folder")
            
            if not all([folder1, folder2, output_folder]):
                task.mark_failed("Missing required folder paths")
                return task
            
            stats = await self._merge_folders(folder1, folder2, output_folder)
            stats.duration_seconds = self.get_elapsed_seconds()
            
            task.metadata["stats"] = stats.to_dict()
            task.mark_completed(stats.to_dict())
            
            self.log_info(f"Merge completed: {stats.successful} files copied")
            return task
            
        except Exception as e:
            self.log_error(f"Merge error: {e}")
            task.mark_failed(str(e))
            return task
    
    async def _merge_folders(
        self, 
        folder1: str, 
        folder2: str, 
        output_folder: str
    ) -> GenerationStats:
        """
        Merge two folders into output folder.
        
        Args:
            folder1: First source folder
            folder2: Second source folder
            output_folder: Destination folder
            
        Returns:
            GenerationStats with file counts
        """
        import shutil
        
        stats = GenerationStats()
        os.makedirs(output_folder, exist_ok=True)
        
        for root_folder in [folder1, folder2]:
            if not os.path.exists(root_folder):
                self.log_warning(f"Source folder not found: {root_folder}")
                continue
            
            for root, _, files in os.walk(root_folder):
                rel_path = os.path.relpath(root, root_folder)
                dest_subfolder = os.path.join(output_folder, rel_path)
                os.makedirs(dest_subfolder, exist_ok=True)
                
                for file in files:
                    ext = os.path.splitext(file)[1].lower()
                    if ext in self.ALLOWED_EXTENSIONS:
                        stats.total_records += 1
                        try:
                            src_path = os.path.join(root, file)
                            dest_path = self._safe_copy(src_path, dest_subfolder)
                            stats.add_success(dest_path)
                        except Exception as e:
                            stats.add_failure(f"Failed to copy {file}: {str(e)}")
        
        return stats
    
    def _safe_copy(self, src_path: str, dest_folder: str) -> str:
        """
        Copy file with automatic renaming if destination exists.
        
        Args:
            src_path: Source file path
            dest_folder: Destination folder
            
        Returns:
            Path to copied file
        """
        import shutil
        
        base, ext = os.path.splitext(os.path.basename(src_path))
        counter = 1
        final_path = os.path.join(dest_folder, base + ext)
        
        while os.path.exists(final_path):
            final_path = os.path.join(dest_folder, f"{base}_{counter}{ext}")
            counter += 1
        
        shutil.copy2(src_path, final_path)
        return final_path
