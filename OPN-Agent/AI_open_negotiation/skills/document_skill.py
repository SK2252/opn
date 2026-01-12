from semantic_kernel.functions import kernel_function
from AI_open_negotiation.utils.logger import log_info, log_error

class DocumentSkill:

    @kernel_function(
        name="create_documents",
        description="Generate Open Negotiation documents"
    )
    def create_documents(self, document_config: dict) -> dict:
        log_info("[DocumentSkill] Document creation started")
        try:
            import os
            import shutil
            import pandas as pd
            from docx import Document
            from docx2pdf import convert
            import pythoncom

            # ================= HELPER FUNCTIONS =================
            def add_dollar_sign(x):
                if pd.isna(x) or str(x).strip().upper() == 'N/A':
                    return 'N/A'
                try:
                    value = float(str(x).replace('$', '').replace(',', '').strip())
                    return f"${value:,.2f}"
                except Exception:
                    return str(x)

            def replace_placeholders(paragraph, replacements, bold_keys):
                full_text = "".join(run.text for run in paragraph.runs)
                replaced = False
                for placeholder, replacement in replacements.items():
                    if placeholder in full_text:
                        full_text = full_text.replace(placeholder, replacement)
                        replaced = True
                if replaced:
                    for run in paragraph.runs:
                        run.text = ""
                    i = 0
                    while i < len(full_text):
                        match_found = False
                        for placeholder, replacement in replacements.items():
                            rep_len = len(replacement)
                            if full_text[i:i + rep_len] == replacement:
                                run = paragraph.add_run(replacement)
                                if placeholder in bold_keys:
                                    run.bold = True
                                i += rep_len
                                match_found = True
                                break
                        if not match_found:
                            paragraph.add_run(full_text[i])
                            i += 1

            def is_allowed_file(filename):
                return os.path.splitext(filename)[1].lower() in {".xls", ".xlsx", ".doc", ".docx", ".pdf"}

            def safe_copy(src_path, dest_path):
                base, ext = os.path.splitext(os.path.basename(src_path))
                counter = 1
                final_path = os.path.join(dest_path, base + ext)
                while os.path.exists(final_path):
                    final_path = os.path.join(dest_path, f"{base}_{counter}{ext}")
                    counter += 1
                shutil.copy2(src_path, final_path)

            def merge_folders(folder1, folder2, output_folder):
                os.makedirs(output_folder, exist_ok=True)
                for root_folder in [folder1, folder2]:
                    for root, _, files in os.walk(root_folder):
                        rel_path = os.path.relpath(root, root_folder)
                        dest_subfolder = os.path.join(output_folder, rel_path)
                        os.makedirs(dest_subfolder, exist_ok=True)
                        for file in files:
                            if is_allowed_file(file):
                                safe_copy(os.path.join(root, file), dest_subfolder)

            # ================= OPEN NEG GROUP =================
            def generate_open_neg_group(input_excel_path, output_group_folder):
                df = pd.read_excel(input_excel_path)
                df.columns = df.columns.str.strip()
                os.makedirs(output_group_folder, exist_ok=True)
                unique_combinations = df[['ProvOrgNPI', 'Provider', 'InsurancePlanName']].drop_duplicates()
                selected_columns = {
                    'CPT_Description': 'Description of item(s) and/or service(s)',
                    'Claim Number': 'Claim Number',
                    'ProvOrgNPI': 'Name of provider, facility, or provider of air ambulance services, and National Provider Identifier (NPI)',
                    'Date of item(s) or service(s)': 'Date provided',
                    'Service code(s)': 'Service code',
                    'Initial Payment': 'Initial payment (if no initial payment amount, write N/A)',
                    'Offer': 'Offer for total out-of- network rate (including any cost sharing)'
                }
                for _, row in unique_combinations.iterrows():
                    filtered_df = df[
                        (df['ProvOrgNPI'] == row['ProvOrgNPI']) &
                        (df['Provider'] == row['Provider']) &
                        (df['InsurancePlanName'] == row['InsurancePlanName'])
                    ]
                    if filtered_df.empty or pd.isna(filtered_df.iloc[0].get('OpenNegGroup')):
                        continue
                    output_df = filtered_df[list(selected_columns.keys())].rename(columns=selected_columns)
                    output_df['Date provided'] = pd.to_datetime(output_df['Date provided'], errors='coerce').dt.strftime('%b %d, %Y')
                    output_df.insert(0, 'SNO', range(1, len(output_df) + 1))
                    output_df['Initial payment (if no initial payment amount, write N/A)'] = output_df['Initial payment (if no initial payment amount, write N/A)'].apply(add_dollar_sign)
                    output_df['Offer for total out-of- network rate (including any cost sharing)'] = output_df['Offer for total out-of- network rate (including any cost sharing)'].apply(add_dollar_sign)
                    safe_npi = "".join(c if c.isalnum() else "_" for c in str(row['ProvOrgNPI']))
                    insurance = str(row['InsurancePlanName']).strip()
                    output_path = os.path.join(output_group_folder, safe_npi, insurance)
                    os.makedirs(output_path, exist_ok=True)
                    group_filename = str(filtered_df.iloc[0]['OpenNegGroup']).strip()
                    if not group_filename.lower().endswith('.xlsx'):
                        group_filename += '.xlsx'
                    output_df.to_excel(os.path.join(output_path, group_filename), index=False)

            # ================= OPEN NEG NOTICE =================
            def generate_open_neg_notice(input_excel_path, template_docx_path, output_notice_folder):
                df = pd.read_excel(input_excel_path, engine='openpyxl')
                df.columns = df.columns.str.strip()
                df = df.drop_duplicates(subset=['ProvOrgNPI', 'Hospital Name', 'OpenNegNotice', 'InsurancePlanName'])
                os.makedirs(output_notice_folder, exist_ok=True)
                for _, row in df.iterrows():
                    if pd.isna(row.get('OpenNegNotice')):
                        continue
                    doc = Document(template_docx_path)
                    replacements = {
                        '{Hospital Name}': str(row['Hospital Name']),
                        '{Provider}': str(row['Provider']),
                        '{InsurancePlanName}': str(row['InsurancePlanName']),
                        '{Notice Date}': str(row['Notice Date']),
                        '{CMS Date1}': str(row['CMS Date1']),
                        '{CMS Date2}': str(row['CMS Date2'])
                    }
                    bold_keys = replacements.keys()
                    for paragraph in doc.paragraphs:
                        replace_placeholders(paragraph, replacements, bold_keys)
                    for table in doc.tables:
                        for row_cells in table.rows:
                            for cell in row_cells.cells:
                                for paragraph in cell.paragraphs:
                                    replace_placeholders(paragraph, replacements, bold_keys)
                    save_path = os.path.join(output_notice_folder, str(row['ProvOrgNPI']), str(row['InsurancePlanName']))
                    os.makedirs(save_path, exist_ok=True)
                    notice_filename = str(row['OpenNegNotice']).strip()
                    base_filename = os.path.splitext(notice_filename)[0]
                    docx_path = os.path.join(save_path, base_filename + ".docx")
                    pdf_path = os.path.join(save_path, base_filename + ".pdf")
                    doc.save(docx_path)
                    convert(docx_path, pdf_path)

            # ================= PIPELINE =================
            pythoncom.CoInitialize()
            try:
                generate_open_neg_group(document_config["excel_path"], document_config["output_group_folder"])
                generate_open_neg_notice(document_config["excel_path"], document_config["template_docx"], document_config["output_notice_folder"])
                merge_folders(document_config["output_group_folder"], document_config["output_notice_folder"], document_config["merged_output_folder"])
            finally:
                pythoncom.CoUninitialize()

            print("[DocumentSkill] Document creation completed successfully")
            log_info(f"[DocumentSkill] Documents created at {document_config['merged_output_folder']}")
            return {
                "status": "SUCCESS",
                "output_folder": document_config["merged_output_folder"]
            }
        
        except Exception as e:
            log_error(f"[DocumentSkill] Error: {e}")
            return {"status": "FAILED", "error": str(e)}